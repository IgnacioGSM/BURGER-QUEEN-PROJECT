
import docx 
from fpdf import FPDF 



document = Document()

# titulo boleta
document.add_heading('Boleta Restaurante', 0)

# agrega informacsiaon
document.add_paragraph('Razón Social del Negocio')
document.add_paragraph('RUT: 12345678-9')
document.add_paragraph('Dirección: Calle Falsa 123')
document.add_paragraph('Teléfono: +56 9 1234 5678')


table = document.add_table(rows=1, cols=4)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Nombre'
hdr_cells[1].text = 'Cantidad'
hdr_cells[2].text = 'Precio Unitario'
hdr_cells[3].text = 'Subtotal'






